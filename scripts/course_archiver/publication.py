from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

from PIL import Image, ImageStat
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.image.image import Image as DocxImage
import pymupdf
from pypdf import PdfReader

from .models import ArticleRecord
from .storage import ArchiveStore
from .utils import atomic_write_json, now_text, safe_filename, sha256_bytes


WORD_EXPORT_SCRIPT = Path(r"C:\Users\Administrator\.codex\skills\windows-office-files\scripts\export_word_fixed.ps1")
XPS_RENDER_SCRIPT = Path(r"C:\Users\Administrator\.codex\skills\windows-office-files\scripts\render_xps_pages.ps1")
PWSH = Path(r"C:\Program Files\PowerShell-7.6.0\pwsh.exe")


class PublicationBuilder:
    def __init__(self, store: ArchiveStore, project_root: Path):
        self.store = store
        self.project_root = Path(project_root).resolve()
        self.tmp_dir = self.project_root / "tmp" / "docs"
        self.tmp_dir.mkdir(parents=True, exist_ok=True)
        self.word_compat_conversions: set[str] = set()
        self._encrypted_backup_root: Optional[Path] = None
        self._encrypted_backup_entries: List[Dict[str, object]] = []

    def build(self, changed_ids: Optional[Sequence[str]] = None, export_fixed: bool = True) -> Dict[str, object]:
        records = [item for item in self.store.load_manifest() if item.status in {"success", "image_only"}]
        grouped: Dict[str, List[ArticleRecord]] = defaultdict(list)
        for record in records:
            grouped[self.volume_key(record)].append(record)
        affected_keys = set(grouped)
        if changed_ids is not None:
            changed = set(changed_ids)
            affected_keys = {self.volume_key(record) for record in records if record.article_id in changed}
        rebuilt_docx: List[Path] = []
        qa_results: List[Dict[str, object]] = []
        for key in sorted(affected_keys, key=self._volume_sort_key):
            volume_records = sorted(grouped[key], key=lambda item: (item.order, item.article_id))
            output_dir = self.store.editions_dir / key
            output_dir.mkdir(parents=True, exist_ok=True)
            filename = f"{safe_filename(self.store.course_title)}_{safe_filename(key)}.docx"
            docx_path = output_dir / filename
            self._build_volume_docx(key, volume_records, docx_path)
            rebuilt_docx.append(docx_path)
            if export_fixed:
                qa_results.append(self._export_and_validate(docx_path, expected_links=len(volume_records)))

        directory_docx = self.store.editions_dir / "总目录.docx"
        self._build_directory_docx(grouped, directory_docx)
        rebuilt_docx.append(directory_docx)
        if export_fixed:
            qa_results.append(self._export_and_validate(directory_docx, expected_links=len(grouped)))

        expected_docx = [
            self.store.editions_dir
            / key
            / f"{safe_filename(self.store.course_title)}_{safe_filename(key)}.docx"
            for key in sorted(grouped, key=self._volume_sort_key)
        ]
        expected_docx.append(directory_docx)
        if changed_ids is not None and export_fixed:
            qa_path = self.store.qa_dir / "publication_qa.json"
            if qa_path.exists():
                try:
                    previous = json.loads(qa_path.read_text(encoding="utf-8"))
                    merged = {
                        str(item.get("docx")): item
                        for item in previous.get("qa", [])
                        if isinstance(item, dict) and item.get("docx")
                    }
                    merged.update({str(item.get("docx")): item for item in qa_results})
                    qa_results = [merged[str(path)] for path in expected_docx if str(path) in merged]
                except (OSError, ValueError, TypeError):
                    pass
        result = {
            "status": "SUCCESS",
            "generated_at": now_text(),
            "volume_count": len(grouped),
            "rebuilt_volume_count": len(affected_keys),
            "docx_files": [str(path) for path in expected_docx],
            "rebuilt_docx_files": [str(path) for path in rebuilt_docx],
            "qa": qa_results,
            "word_compat_converted_images": len(self.word_compat_conversions),
        }
        atomic_write_json(self.store.qa_dir / "publication_qa.json", result)
        return result

    @staticmethod
    def volume_key(record: ArticleRecord) -> str:
        if any(label in (record.section or "") for label in ("发刊词", "课前加餐")):
            return "卷00_导读与课前加餐"
        for text in (record.section, record.published_at, record.title):
            match = re.search(r"(20\d{2})\s*年\s*(\d{1,2})\s*月", text or "")
            if match:
                return f"{int(match.group(1)):04d}-{int(match.group(2)):02d}"
        return "卷00_导读与课前加餐"

    @staticmethod
    def _volume_sort_key(value: str) -> Tuple[int, str]:
        return (0, value) if value.startswith("卷00") else (1, value)

    def _new_document(self) -> Document:
        document = Document()
        section = document.sections[0]
        section.page_width = Mm(170)
        section.page_height = Mm(240)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(18)
        section.left_margin = Mm(22)
        section.right_margin = Mm(18)
        section.header_distance = Mm(9)
        section.footer_distance = Mm(9)
        section.different_first_page_header_footer = True
        self._enable_mirror_margins(document)
        self._enable_even_odd_headers(document)
        self._configure_styles(document)
        return document

    def _configure_styles(self, document: Document) -> None:
        normal = document.styles["Normal"]
        self._set_style_font(normal, "宋体", 10.5)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.first_line_indent = Pt(21)

        for style_name, font_name, size, color in (
            ("Title", "黑体", 26, RGBColor(34, 34, 34)),
            ("Subtitle", "微软雅黑", 12, RGBColor(90, 90, 90)),
            ("Heading 1", "黑体", 18, RGBColor(28, 28, 28)),
            ("Heading 2", "黑体", 14, RGBColor(45, 45, 45)),
            ("Heading 3", "黑体", 12, RGBColor(65, 65, 65)),
        ):
            style = document.styles[style_name]
            self._set_style_font(style, font_name, size)
            style.font.color.rgb = color
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.page_break_before = False
        document.styles["Heading 1"].paragraph_format.space_before = Pt(14)
        document.styles["Heading 1"].paragraph_format.space_after = Pt(10)
        document.styles["Heading 2"].paragraph_format.space_before = Pt(10)
        document.styles["Heading 2"].paragraph_format.space_after = Pt(6)

        if "Caption CN" not in [style.name for style in document.styles]:
            caption = document.styles.add_style("Caption CN", WD_STYLE_TYPE.PARAGRAPH)
        else:
            caption = document.styles["Caption CN"]
        self._set_style_font(caption, "微软雅黑", 9)
        caption.font.color.rgb = RGBColor(100, 100, 100)
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)

        if "Metadata" not in [style.name for style in document.styles]:
            metadata = document.styles.add_style("Metadata", WD_STYLE_TYPE.PARAGRAPH)
        else:
            metadata = document.styles["Metadata"]
        self._set_style_font(metadata, "微软雅黑", 9)
        metadata.font.color.rgb = RGBColor(105, 105, 105)
        metadata.paragraph_format.first_line_indent = Pt(0)
        metadata.paragraph_format.space_after = Pt(8)

    @staticmethod
    def _set_style_font(style, name: str, size: float) -> None:
        style.font.name = name
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        style._element.rPr.rFonts.set(qn("w:ascii"), name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), name)

    def _build_volume_docx(self, volume_key: str, records: Sequence[ArticleRecord], output_path: Path) -> None:
        document = self._new_document()
        self._add_cover(document, volume_key, len(records))
        document.add_page_break()
        self._add_notice(document)
        document.add_page_break()
        heading = document.add_heading("目录", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_paragraph = document.add_paragraph()
        self._add_toc_field(toc_paragraph)

        for index, record in enumerate(records):
            title_paragraph = document.add_heading(record.title, level=1)
            if index:
                title_paragraph.paragraph_format.page_break_before = True
            title_paragraph.paragraph_format.keep_with_next = True
            metadata = document.add_paragraph(style="Metadata")
            metadata.add_run(
                f"第 {record.order} 讲　{record.section or '未分组'}　"
                f"{record.published_at or '页面未标明发布日期'}　{self._type_label(record.content_type)}"
            )
            article_path = self.store.root / Path(record.markdown_path)
            body = article_path.read_text(encoding="utf-8") if article_path.exists() else ""
            body = self._strip_archive_header(body)
            self._append_markdown(document, body, article_path.parent, record)

        self._set_headers_and_footers(document, volume_key)
        self._set_update_fields_on_open(document)
        self._atomic_save_docx(document, output_path)
        self._validate_docx(output_path, expected_articles=len(records), require_linked_toc=True)

    def _build_directory_docx(self, grouped: Dict[str, List[ArticleRecord]], output_path: Path) -> None:
        document = self._new_document()
        self._add_cover(document, "总目录", sum(len(items) for items in grouped.values()))
        document.add_page_break()
        document.add_heading("分卷目录", level=1)
        for key in sorted(grouped, key=self._volume_sort_key):
            records = sorted(grouped[key], key=lambda item: (item.order, item.article_id))
            document.add_heading(f"{key}（{len(records)} 讲）", level=2)
            target_dir = self.store.editions_dir / key
            target_name = f"{safe_filename(self.store.course_title)}_{safe_filename(key)}.pdf"
            paragraph = document.add_paragraph()
            self._add_hyperlink(paragraph, f"打开 {key} PDF", str((target_dir / target_name).resolve()))
            for record in records:
                item = document.add_paragraph(style="List Bullet")
                item.paragraph_format.first_line_indent = Pt(0)
                item.add_run(f"{record.order:03d}　{record.title}")
        self._set_headers_and_footers(document, "总目录")
        self._set_update_fields_on_open(document)
        self._atomic_save_docx(document, output_path)
        self._validate_docx(output_path, expected_articles=0, expected_external_links=len(grouped))

    def _add_cover(self, document: Document, volume_key: str, count: int) -> None:
        for _ in range(4):
            document.add_paragraph()
        title = document.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run(self.store.course_title)
        subtitle = document.add_paragraph(style="Subtitle")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f"{volume_key} · {count} 讲")
        date_paragraph = document.add_paragraph(style="Subtitle")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run(f"整理日期：{datetime.now():%Y年%m月%d日}")

    def _add_notice(self, document: Document) -> None:
        document.add_heading("阅读说明", level=1)
        for text in (
            "本文件由用户本人有权访问的课程正文整理而成，仅供个人离线学习与阅读，不用于公开传播或商业发行。",
            "正文、标题和插图保持课程原有顺序；音频、视频、评论区、推荐内容和界面素材未收入本卷。",
            "课程持续更新，本卷版本以封面所示整理日期为准。",
        ):
            document.add_paragraph(text)

    @staticmethod
    def _type_label(content_type: str) -> str:
        return {
            "article": "正文",
            "live_replay": "直播回放文字稿",
            "live_qa": "直播问答文字稿",
            "image_only": "图片型正文",
            "article_with_images": "图文正文",
        }.get(content_type, content_type)

    @staticmethod
    def _strip_archive_header(text: str) -> str:
        lines = text.splitlines()
        index = 0
        if lines and lines[0].startswith("# "):
            index = 1
        while index < len(lines) and (not lines[index].strip() or lines[index].lstrip().startswith(">")):
            index += 1
        return "\n".join(lines[index:]).strip()

    def _append_markdown(self, document: Document, markdown: str, base_dir: Path, record: ArticleRecord) -> None:
        paragraph_buffer: List[str] = []

        def flush() -> None:
            if not paragraph_buffer:
                return
            text = " ".join(item.strip() for item in paragraph_buffer if item.strip()).strip()
            paragraph_buffer.clear()
            if text:
                paragraph = document.add_paragraph(text)
                paragraph.paragraph_format.widow_control = True

        image_index = 0
        for raw_line in markdown.splitlines():
            line = raw_line.rstrip()
            image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
            if image_match:
                flush()
                image_index += 1
                image_path = (base_dir / image_match.group(2)).resolve()
                if image_path.exists():
                    self._add_image(document, image_path, image_match.group(1) or f"正文插图 {image_index}", record, image_index)
                else:
                    document.add_paragraph(f"[图片文件缺失：{image_path.name}]", style="Metadata")
                continue
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                flush()
                level = min(3, max(2, len(heading_match.group(1))))
                document.add_heading(self._inline_plain(heading_match.group(2)), level=level)
                continue
            if re.match(r"^[-*+]\s+", line):
                flush()
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.add_run(self._inline_plain(re.sub(r"^[-*+]\s+", "", line)))
                continue
            if re.match(r"^\d+[.)]\s+", line):
                flush()
                paragraph = document.add_paragraph(style="List Number")
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.add_run(self._inline_plain(re.sub(r"^\d+[.)]\s+", "", line)))
                continue
            if line.startswith(">"):
                flush()
                paragraph = document.add_paragraph(self._inline_plain(line.lstrip("> ")), style="Quote")
                paragraph.paragraph_format.first_line_indent = Pt(0)
                continue
            if not line.strip():
                flush()
                continue
            paragraph_buffer.append(self._inline_plain(line))
        flush()

    @staticmethod
    def _inline_plain(text: str) -> str:
        text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
        text = re.sub(r"[*_`]+", "", text)
        return text.strip()

    def _add_image(self, document: Document, image_path: Path, alt: str, record: ArticleRecord, image_index: int) -> None:
        content_width_mm = 130.0
        max_height_mm = 174.0
        segments = self._image_segments(image_path, content_width_mm, max_height_mm)
        for segment_index, segment in enumerate(segments, 1):
            word_image = self._word_compatible_image(segment)
            with Image.open(word_image) as image:
                width_px, height_px = image.size
            width_mm = content_width_mm
            height_mm = width_mm * height_px / max(width_px, 1)
            if height_mm > max_height_mm:
                height_mm = max_height_mm
                width_mm = height_mm * width_px / max(height_px, 1)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_together = True
            run = paragraph.add_run()
            run.add_picture(str(word_image), width=Mm(width_mm), height=Mm(height_mm))
            caption_text = f"图 {record.order}-{image_index}　{alt}"
            if len(segments) > 1:
                caption_text += f"（{segment_index}/{len(segments)}）"
            document.add_paragraph(caption_text, style="Caption CN")

    def _image_segments(self, image_path: Path, content_width_mm: float, max_height_mm: float) -> List[Path]:
        with Image.open(image_path) as image:
            width, height = image.size
            max_segment_height = max(1, int(width * max_height_mm / content_width_mm))
            if height <= max_segment_height * 1.05:
                return [image_path]
            segment_dir = self.tmp_dir / "image_segments" / safe_filename(image_path.stem, 50)
            segment_dir.mkdir(parents=True, exist_ok=True)
            outputs: List[Path] = []
            top = 0
            index = 1
            while top < height:
                bottom = min(height, top + max_segment_height)
                segment = image.crop((0, top, width, bottom))
                output = segment_dir / f"{index:03d}.png"
                segment.save(output, format="PNG", optimize=True)
                outputs.append(output)
                top = bottom
                index += 1
            return outputs

    def _word_compatible_image(self, image_path: Path) -> Path:
        try:
            DocxImage.from_file(str(image_path))
            return image_path
        except Exception:
            digest = sha256_bytes(image_path.read_bytes())[:20]
            output_dir = self.tmp_dir / "word_compatible"
            output_dir.mkdir(parents=True, exist_ok=True)
            output = output_dir / f"{digest}.png"
            if not output.exists():
                temporary = output.with_name(f".{output.name}.{os.getpid()}.tmp")
                with Image.open(image_path) as source:
                    converted = source.convert("RGBA") if "A" in source.getbands() else source.convert("RGB")
                    converted.save(temporary, format="PNG", optimize=True)
                os.replace(temporary, output)
            DocxImage.from_file(str(output))
            self.word_compat_conversions.add(str(image_path))
            return output

    def _set_headers_and_footers(self, document: Document, volume_key: str) -> None:
        for section in document.sections:
            header = section.header
            header.is_linked_to_previous = False
            paragraph = header.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.text = self.store.course_title
            self._format_runs(paragraph, "微软雅黑", 8.5, RGBColor(115, 115, 115))
            even_header = section.even_page_header
            even_header.is_linked_to_previous = False
            even_paragraph = even_header.paragraphs[0]
            even_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            even_paragraph.text = volume_key
            self._format_runs(even_paragraph, "微软雅黑", 8.5, RGBColor(115, 115, 115))
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_field(footer_paragraph, "PAGE")

    @staticmethod
    def _format_runs(paragraph, font_name: str, size: float, color: RGBColor) -> None:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    @staticmethod
    def _add_field(paragraph, instruction: str) -> None:
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.extend([fld_begin, instr, fld_separate, fld_end])

    def _add_toc_field(self, paragraph) -> None:
        self._add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')

    @staticmethod
    def _add_hyperlink(paragraph, text: str, target: str) -> None:
        part = paragraph.part
        relationship_id = part.relate_to(target, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)
        run = OxmlElement("w:r")
        properties = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        properties.extend([color, underline])
        run.append(properties)
        node = OxmlElement("w:t")
        node.text = text
        run.append(node)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    @staticmethod
    def _enable_mirror_margins(document: Document) -> None:
        settings = document.settings._element
        if settings.find(qn("w:mirrorMargins")) is None:
            settings.append(OxmlElement("w:mirrorMargins"))

    @staticmethod
    def _enable_even_odd_headers(document: Document) -> None:
        settings = document.settings._element
        if settings.find(qn("w:evenAndOddHeaders")) is None:
            settings.append(OxmlElement("w:evenAndOddHeaders"))

    @staticmethod
    def _set_update_fields_on_open(document: Document) -> None:
        settings = document.settings._element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    @staticmethod
    def _atomic_save_docx(document: Document, output_path: Path) -> None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        temporary = output_path.with_name(f".{output_path.stem}.{os.getpid()}.tmp.docx")
        document.save(temporary)
        os.replace(temporary, output_path)

    @staticmethod
    def _validate_docx(
        path: Path,
        expected_articles: int,
        require_linked_toc: bool = False,
        expected_external_links: int = 0,
    ) -> None:
        if path.read_bytes()[:2] != b"PK":
            raise ValueError(f"DOCX 签名无效: {path}")
        if not zipfile.is_zipfile(path):
            raise ValueError(f"DOCX ZIP 结构无效: {path}")
        document = Document(path)
        headings = [paragraph for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading")]
        if expected_articles and len(headings) < expected_articles:
            raise ValueError(f"DOCX 标题数不足: {len(headings)} < {expected_articles}")
        if require_linked_toc:
            instructions = "".join(
                node.text or "" for node in document._element.iter(qn("w:instrText"))
            )
            if "TOC" not in instructions or "\\h" not in instructions:
                raise ValueError(f"DOCX 目录缺少内部超链接字段: {path}")
        external_links = sum(
            1
            for relationship in document.part.rels.values()
            if relationship.reltype.endswith("/hyperlink") and relationship.is_external
        )
        if external_links < expected_external_links:
            raise ValueError(
                f"DOCX 外部目录链接数不足: {external_links} < {expected_external_links}, file={path.name}"
            )

    def _export_and_validate(self, docx_path: Path, expected_links: int = 0) -> Dict[str, object]:
        if not WORD_EXPORT_SCRIPT.exists() or not XPS_RENDER_SCRIPT.exists() or not PWSH.exists():
            raise RuntimeError("Windows Office 导出脚本或 PowerShell 7 不可用")
        original_hash = sha256_bytes(docx_path.read_bytes())
        pdf_path = docx_path.with_suffix(".pdf")
        xps_path = self.store.qa_dir / "xps" / f"{docx_path.stem}.xps"
        pages_dir = self.store.qa_dir / "pages" / docx_path.stem
        xps_path.parent.mkdir(parents=True, exist_ok=True)
        pages_dir.mkdir(parents=True, exist_ok=True)
        # 页面图是可再生 QA 缓存。分卷页数变化时，遗留尾页会造成错误的
        # 页数/空白页判断，因此在本次 XPS 渲染前只清除该分卷的旧检查图。
        for stale_page in pages_dir.glob("page-*.png"):
            stale_page.unlink()
        xps_result = self._run_json_command([
            str(PWSH), "-NoProfile", "-File", str(WORD_EXPORT_SCRIPT),
            "-InputPath", str(docx_path), "-OutputPath", str(xps_path), "-Format", "Xps", "-Force",
        ])
        render_result = self._run_json_command([
            str(PWSH), "-NoProfile", "-File", str(XPS_RENDER_SCRIPT),
            "-InputPath", str(xps_path), "-OutputDirectory", str(pages_dir), "-Dpi", "144", "-Force",
        ])
        if sha256_bytes(docx_path.read_bytes()) != original_hash:
            raise ValueError(f"Word 导出后 DOCX 哈希发生变化: {docx_path}")
        xps_pages = int(render_result.get("pages", 0))
        word_pages = int(xps_result.get("pages", 0))
        if xps_pages != word_pages:
            raise ValueError(
                f"页数不一致: Word={word_pages}, XPS={xps_pages}, file={docx_path.name}"
            )
        # PowerShell launched through an interactive Windows shell may encode its JSON
        # path strings with the active console code page. Enumerate the known output
        # directory instead of trusting those display strings; numeric QA fields remain
        # safe to parse from JSON.
        page_files = sorted(pages_dir.glob("page-*.png"))
        if len(page_files) != xps_pages:
            raise ValueError(
                f"XPS 渲染页文件数不一致: files={len(page_files)}, XPS={xps_pages}, file={docx_path.name}"
            )
        blank_pages = self._detect_blank_pages(page_files)
        if blank_pages:
            raise ValueError(f"检测到疑似空白页: {blank_pages}, file={docx_path.name}")
        pdf_result = self._convert_xps_to_native_pdf(
            xps_path,
            pdf_path,
            word_pages,
            expected_links=expected_links,
        )
        pdf_blank_pages = self._detect_blank_pdf_pages(pdf_path)
        if pdf_blank_pages:
            raise ValueError(f"明文 PDF 检测到疑似空白页: {pdf_blank_pages}, file={docx_path.name}")
        return {
            "docx": str(docx_path),
            "pdf": str(pdf_path),
            "xps": str(xps_path),
            "pages": word_pages,
            "pdf_signature": pdf_result["pdf_signature"],
            "pdf_encrypted": pdf_result["pdf_encrypted"],
            "pdf_source": "xps",
            "pdf_sha256": pdf_result["pdf_sha256"],
            "rendered_pages": len(page_files),
            "pdf_rendered_pages": pdf_result["pages"],
            "pdf_links": pdf_result["pdf_links"],
            "blank_pages": blank_pages,
            "docx_sha256": original_hash,
        }

    def _convert_xps_to_native_pdf(
        self,
        xps_path: Path,
        pdf_path: Path,
        expected_pages: int,
        expected_links: int = 0,
    ) -> Dict[str, object]:
        with xps_path.open("rb") as stream:
            if stream.read(2) != b"PK":
                raise ValueError(f"XPS ZIP 签名无效: {xps_path}")
        pdf_tmp_dir = self.project_root / "tmp" / "pdfs"
        pdf_tmp_dir.mkdir(parents=True, exist_ok=True)
        handle, temporary_name = tempfile.mkstemp(prefix=f".{safe_filename(xps_path.stem)}.", suffix=".pdf", dir=pdf_tmp_dir)
        os.close(handle)
        temporary_pdf = Path(temporary_name)
        try:
            with pymupdf.open(str(xps_path)) as xps_document:
                if xps_document.page_count != expected_pages:
                    raise ValueError(
                        f"XPS 页数不一致: expected={expected_pages}, actual={xps_document.page_count}, file={xps_path.name}"
                    )
                temporary_pdf.write_bytes(xps_document.convert_to_pdf())
            self._validate_native_pdf(temporary_pdf, expected_pages, expected_links=expected_links)
            blank_pages = self._detect_blank_pdf_pages(temporary_pdf)
            if blank_pages:
                raise ValueError(f"明文 PDF 检测到疑似空白页: {blank_pages}, file={xps_path.name}")
            return self._promote_native_pdf(
                temporary_pdf,
                pdf_path,
                expected_pages,
                expected_links=expected_links,
            )
        finally:
            if temporary_pdf.exists():
                temporary_pdf.unlink()

    @staticmethod
    def _validate_native_pdf(
        pdf_path: Path,
        expected_pages: int,
        expected_links: int = 0,
    ) -> Dict[str, object]:
        with pdf_path.open("rb") as stream:
            signature = stream.read(16)
        if signature.startswith(b"%TSD-Header-"):
            raise ValueError(f"PDF 仍被透明加密: {pdf_path}")
        if not signature.startswith(b"%PDF-"):
            raise ValueError(f"PDF 签名无效: {pdf_path}")
        reader = PdfReader(str(pdf_path))
        if reader.is_encrypted:
            raise ValueError(f"PDF 包含密码或权限加密: {pdf_path}")
        pages = len(reader.pages)
        if pages != expected_pages:
            raise ValueError(f"PDF 页数不一致: expected={expected_pages}, actual={pages}, file={pdf_path.name}")
        with pymupdf.open(str(pdf_path)) as document:
            pdf_links = sum(len(page.get_links()) for page in document)
        if pdf_links < expected_links:
            raise ValueError(
                f"PDF 目录链接数不足: {pdf_links} < {expected_links}, file={pdf_path.name}"
            )
        return {
            "pages": pages,
            "pdf_signature": "NATIVE_PDF",
            "pdf_encrypted": False,
            "pdf_links": pdf_links,
            "pdf_sha256": sha256_bytes(pdf_path.read_bytes()),
        }

    def _promote_native_pdf(
        self,
        temporary_pdf: Path,
        pdf_path: Path,
        expected_pages: int,
        expected_links: int = 0,
    ) -> Dict[str, object]:
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        backup_path: Optional[Path] = None
        rollback_path: Optional[Path] = None
        if pdf_path.exists():
            with pdf_path.open("rb") as stream:
                current_signature = stream.read(16)
            if current_signature.startswith(b"%TSD-Header-"):
                backup_path = self._backup_encrypted_pdf(pdf_path)
            elif current_signature.startswith(b"%PDF-"):
                rollback_dir = self.project_root / "tmp" / "pdfs" / "rollback"
                rollback_dir.mkdir(parents=True, exist_ok=True)
                rollback_path = rollback_dir / f"{safe_filename(pdf_path.stem)}.{os.getpid()}.pdf"
                if rollback_path.exists():
                    raise FileExistsError(f"PDF 回滚文件已存在: {rollback_path}")
                os.replace(pdf_path, rollback_path)
            else:
                raise ValueError(f"拒绝替换未识别的 PDF 文件: {pdf_path}")
        try:
            os.replace(temporary_pdf, pdf_path)
            result = self._validate_native_pdf(pdf_path, expected_pages, expected_links=expected_links)
        except Exception:
            if pdf_path.exists():
                pdf_path.unlink()
            previous_path = backup_path or rollback_path
            if previous_path is not None and previous_path.exists():
                os.replace(previous_path, pdf_path)
            if backup_path is not None:
                self._remove_backup_entry(backup_path)
            raise
        if rollback_path is not None and rollback_path.exists():
            rollback_path.unlink()
        return result

    def _backup_encrypted_pdf(self, pdf_path: Path) -> Path:
        editions_root = self.store.editions_dir.resolve()
        resolved_pdf = pdf_path.resolve()
        relative_path = resolved_pdf.relative_to(editions_root)
        if self._encrypted_backup_root is None:
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            self._encrypted_backup_root = self.store.qa_dir / "encrypted_pdf_backup" / stamp
        backup_path = self._encrypted_backup_root / relative_path
        if backup_path.exists():
            raise FileExistsError(f"加密 PDF 备份已存在: {backup_path}")
        backup_path.parent.mkdir(parents=True, exist_ok=True)
        entry = {
            "original_path": str(resolved_pdf),
            "backup_path": str(backup_path.resolve()),
            "size": pdf_path.stat().st_size,
            "sha256": sha256_bytes(pdf_path.read_bytes()),
            "signature": "TSD_TRANSPARENT_ENCRYPTION",
        }
        os.replace(pdf_path, backup_path)
        self._encrypted_backup_entries.append(entry)
        self._write_backup_manifest()
        return backup_path

    def _remove_backup_entry(self, backup_path: Path) -> None:
        resolved = str(backup_path.resolve())
        self._encrypted_backup_entries = [
            item for item in self._encrypted_backup_entries if item.get("backup_path") != resolved
        ]
        self._write_backup_manifest()

    def _write_backup_manifest(self) -> None:
        if self._encrypted_backup_root is None:
            return
        atomic_write_json(
            self._encrypted_backup_root / "manifest.json",
            {
                "created_at": now_text(),
                "file_count": len(self._encrypted_backup_entries),
                "files": self._encrypted_backup_entries,
            },
        )

    @staticmethod
    def _detect_blank_pdf_pages(pdf_path: Path) -> List[int]:
        blank: List[int] = []
        with pymupdf.open(str(pdf_path)) as document:
            for index, page in enumerate(document, 1):
                pixmap = page.get_pixmap(matrix=pymupdf.Matrix(0.5, 0.5), colorspace=pymupdf.csGRAY, alpha=False)
                image = Image.frombytes("L", (pixmap.width, pixmap.height), pixmap.samples)
                image.thumbnail((400, 600))
                stat = ImageStat.Stat(image)
                dark_pixels = sum(image.histogram()[:245])
                total_pixels = max(1, image.width * image.height)
                if stat.mean[0] > 253.5 and dark_pixels / total_pixels < 0.001:
                    blank.append(index)
        return blank

    @staticmethod
    def _run_json_command(command: Sequence[str]) -> Dict[str, object]:
        completed = subprocess.run(command, check=False, capture_output=True, text=True, encoding="utf-8", errors="replace")
        if completed.returncode != 0:
            raise RuntimeError(f"命令失败 ({completed.returncode}): {' '.join(command)}\n{completed.stderr.strip()}")
        output = completed.stdout.strip().splitlines()
        if not output:
            raise RuntimeError(f"命令没有返回 JSON: {' '.join(command)}")
        return json.loads(output[-1])

    @staticmethod
    def _detect_blank_pages(paths: Sequence[Path]) -> List[int]:
        blank: List[int] = []
        for index, path in enumerate(paths, 1):
            with Image.open(path) as image:
                gray = image.convert("L")
                gray.thumbnail((400, 600))
                stat = ImageStat.Stat(gray)
                dark_pixels = sum(gray.histogram()[:245])
                total_pixels = max(1, gray.width * gray.height)
                if stat.mean[0] > 253.5 and dark_pixels / total_pixels < 0.001:
                    blank.append(index)
        return blank
