from pathlib import Path

import pytest
from docx import Document
from PIL import Image
from pypdf import PdfWriter
from docx.oxml.ns import qn

from course_archiver.models import ArticleRecord, ImageRecord
from course_archiver.publication import PublicationBuilder
from course_archiver.storage import ArchiveStore


def test_build_volume_docx(tmp_path: Path):
    store = ArchiveStore(tmp_path / "output", "何刚·投资参考（年度日更）")
    image_path = store.images_dir / "id1" / "001.png"
    image_path.parent.mkdir(parents=True, exist_ok=True)
    Image.new("RGB", (800, 600), "white").save(image_path)
    record = ArticleRecord(
        order=1,
        article_id="id1",
        url="https://www.dedao.cn/course/article?id=id1",
        title="001｜测试课程",
        section="2026年1月(已更新1讲)",
        content_type="article",
        published_at="2026年1月1日",
        markdown_path="archive/articles/0001_id1.md",
        text_char_count=100,
        body_sha256="abc",
        scrape_time="2026-08-09 09:00:00",
        images=[ImageRecord(1, "https://example.test/1.png", "archive/images/id1/001.png", "image/png", 800, 600, "x", image_path.stat().st_size)],
    )
    store.save_article_markdown(record, "## 小节\n\n这是一段用于排版测试的正文。\n\n![正文插图](../images/id1/001.png)")
    store.save_manifest([record])
    result = PublicationBuilder(store, tmp_path).build(export_fixed=False)
    assert result["volume_count"] == 1
    volume = store.editions_dir / "2026-01" / "何刚·投资参考（年度日更）_2026-01.docx"
    assert volume.exists()
    document = Document(volume)
    assert any("测试课程" in paragraph.text for paragraph in document.paragraphs)
    lecture_title = next(paragraph for paragraph in document.paragraphs if "测试课程" in paragraph.text)
    assert lecture_title.paragraph_format.page_break_before is None
    assert len(document.inline_shapes) == 1
    instructions = "".join(
        node.text or "" for node in document._element.iter(qn("w:instrText"))
    )
    assert "TOC" in instructions
    assert "\\h" in instructions


def test_build_directory_docx_links_to_volume_pdf(tmp_path: Path):
    store = ArchiveStore(tmp_path / "output", "测试课程")
    record = ArticleRecord(
        order=1,
        article_id="id1",
        url="https://www.dedao.cn/course/article?id=id1",
        title="001｜测试课程",
        section="2026年1月",
        content_type="article",
        published_at="2026年1月1日",
        markdown_path="archive/articles/0001_id1.md",
        text_char_count=100,
        body_sha256="abc",
        scrape_time="2026-08-09 09:00:00",
        images=[],
    )
    output_path = store.editions_dir / "总目录.docx"
    PublicationBuilder(store, tmp_path)._build_directory_docx(
        {"2026-01": [record]},
        output_path,
    )

    document = Document(output_path)
    hyperlink_targets = [
        relationship.target_ref
        for relationship in document.part.rels.values()
        if relationship.reltype.endswith("/hyperlink") and relationship.is_external
    ]
    assert len(hyperlink_targets) == 1
    assert hyperlink_targets[0].endswith("测试课程_2026-01.pdf")


def _write_pdf(path: Path, pages: int = 1) -> None:
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=200, height=300)
    with path.open("wb") as stream:
        writer.write(stream)


def test_validate_native_pdf(tmp_path: Path):
    pdf_path = tmp_path / "native.pdf"
    _write_pdf(pdf_path, pages=2)
    result = PublicationBuilder._validate_native_pdf(pdf_path, expected_pages=2)
    assert result["pdf_signature"] == "NATIVE_PDF"
    assert result["pdf_encrypted"] is False
    assert result["pages"] == 2


def test_validate_native_pdf_rejects_tsd(tmp_path: Path):
    pdf_path = tmp_path / "encrypted.pdf"
    pdf_path.write_bytes(b"%TSD-Header-###%encrypted")
    with pytest.raises(ValueError, match="透明加密"):
        PublicationBuilder._validate_native_pdf(pdf_path, expected_pages=1)


def test_validate_native_pdf_rejects_page_mismatch(tmp_path: Path):
    pdf_path = tmp_path / "wrong-pages.pdf"
    _write_pdf(pdf_path, pages=1)
    with pytest.raises(ValueError, match="PDF 页数不一致"):
        PublicationBuilder._validate_native_pdf(pdf_path, expected_pages=2)


def test_validate_native_pdf_rejects_missing_directory_links(tmp_path: Path):
    pdf_path = tmp_path / "no-links.pdf"
    _write_pdf(pdf_path, pages=1)
    with pytest.raises(ValueError, match="PDF 目录链接数不足"):
        PublicationBuilder._validate_native_pdf(pdf_path, expected_pages=1, expected_links=1)


def test_promote_native_pdf_backs_up_tsd(tmp_path: Path):
    store = ArchiveStore(tmp_path / "output", "test-course")
    builder = PublicationBuilder(store, tmp_path)
    pdf_path = store.editions_dir / "volume.pdf"
    pdf_path.write_bytes(b"%TSD-Header-###%encrypted")
    temporary_pdf = tmp_path / "native.tmp.pdf"
    _write_pdf(temporary_pdf)

    result = builder._promote_native_pdf(temporary_pdf, pdf_path, expected_pages=1)

    assert result["pdf_signature"] == "NATIVE_PDF"
    assert pdf_path.read_bytes().startswith(b"%PDF-")
    assert builder._encrypted_backup_root is not None
    manifest_path = builder._encrypted_backup_root / "manifest.json"
    assert manifest_path.exists()
    assert builder._encrypted_backup_entries[0]["signature"] == "TSD_TRANSPARENT_ENCRYPTION"
    assert Path(builder._encrypted_backup_entries[0]["backup_path"]).read_bytes().startswith(b"%TSD-Header-")


def test_conversion_failure_keeps_existing_pdf(tmp_path: Path):
    store = ArchiveStore(tmp_path / "output", "test-course")
    builder = PublicationBuilder(store, tmp_path)
    pdf_path = store.editions_dir / "volume.pdf"
    original = b"%TSD-Header-###%encrypted"
    pdf_path.write_bytes(original)
    invalid_xps = tmp_path / "invalid.xps"
    invalid_xps.write_bytes(b"not-an-xps")

    with pytest.raises(ValueError, match="XPS ZIP 签名无效"):
        builder._convert_xps_to_native_pdf(invalid_xps, pdf_path, expected_pages=1)

    assert pdf_path.read_bytes() == original
    assert builder._encrypted_backup_root is None
